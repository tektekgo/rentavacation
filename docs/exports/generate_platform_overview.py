"""
Generate a branded .docx Platform Overview document for Rent-A-Vacation.
Reuses brand helpers from generate_docx.py.
"""

import os
import sys
from datetime import datetime

# Add exports dir to path so we can import helpers
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

from generate_docx import (
    create_branded_doc,
    add_logo_header,
    add_page_numbers,
    add_metadata,
    add_body,
    add_blockquote,
    add_table_from_data,
    add_horizontal_rule,
    add_footer,
    DEEP_TEAL,
    WARM_CORAL,
    DARK_NAVY,
    BRAND_FONT,
)
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_numbered_list(doc, items, bold_prefix=False):
    """Add a numbered list with brand styling."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{i}. ")
        run.font.name = BRAND_FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DEEP_TEAL

        # Check for "bold → rest" pattern
        if "\u2192" in item:
            parts = item.split("\u2192", 1)
            run2 = p.add_run(parts[0].strip())
            run2.font.name = BRAND_FONT
            run2.font.size = Pt(10)
            run2.font.bold = True
            run2.font.color.rgb = DARK_NAVY
            run3 = p.add_run(f" \u2192 {parts[1].strip()}")
            run3.font.name = BRAND_FONT
            run3.font.size = Pt(10)
            run3.font.color.rgb = DARK_NAVY
        else:
            run2 = p.add_run(item)
            run2.font.name = BRAND_FONT
            run2.font.size = Pt(10)
            run2.font.color.rgb = DARK_NAVY


def add_bullet_list(doc, items):
    """Add a bullet list with brand styling."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

        # Check for "Label: Description" pattern
        if ": " in item:
            label, desc = item.split(": ", 1)
            run_bullet = p.add_run("\u2022 ")
            run_bullet.font.name = BRAND_FONT
            run_bullet.font.size = Pt(10)
            run_bullet.font.color.rgb = DEEP_TEAL
            run_label = p.add_run(f"{label}: ")
            run_label.font.name = BRAND_FONT
            run_label.font.size = Pt(10)
            run_label.font.bold = True
            run_label.font.color.rgb = DARK_NAVY
            run_desc = p.add_run(desc)
            run_desc.font.name = BRAND_FONT
            run_desc.font.size = Pt(10)
            run_desc.font.color.rgb = DARK_NAVY
        else:
            run = p.add_run(f"\u2022 {item}")
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_NAVY


def generate_platform_overview():
    today = datetime.now().strftime("%m%d%Y")
    doc = create_branded_doc("Platform Overview")
    add_logo_header(doc, doc_title="Platform Overview \u2014 What\u2019s Been Built")
    add_page_numbers(doc)

    # Metadata
    add_metadata(doc, [
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Version", "v0.9.0 (Pre-Launch)"),
        ("Website", "https://rent-a-vacation.com"),
        ("Repository", "github.com/rent-a-vacation/rav-website"),
    ])

    # ── What It Is ──
    doc.add_heading("What It Is", level=1)
    add_body(
        doc,
        "A marketplace where timeshare owners can rent out their unused vacation weeks "
        "to travelers, with RAV earning a 15% commission. Think Airbnb, but specifically "
        "for timeshare inventory across Hilton, Marriott, Disney, and 6 other vacation "
        "club brands (117 resorts total).",
    )

    # ── Tech Stack ──
    doc.add_heading("Tech Stack", level=1)
    add_bullet_list(doc, [
        "Frontend: React + TypeScript + Vite + Tailwind + shadcn/ui",
        "Backend: Supabase (PostgreSQL, Auth, Edge Functions, RLS)",
        "Payments: Stripe (checkout, Connect payouts, webhooks)",
        "Voice: VAPI (Deepgram STT + GPT-4o-mini + ElevenLabs TTS)",
        "Text Chat: OpenRouter (RAVIO assistant)",
        "Deployment: Vercel (frontend) + Supabase (backend)",
    ])

    add_horizontal_rule(doc)

    # ── Core User Journeys ──
    doc.add_heading("Core User Journeys", level=1)

    doc.add_heading("Property Owner Flow", level=2)
    add_numbered_list(doc, [
        "Sign up \u2192 pending approval by RAV admin",
        "Add property (9 brands supported) \u2192 create listing with nightly rate",
        "Listing goes to pending_approval \u2192 RAV admin approves/rejects",
        "Once booked \u2192 owner confirms resort reservation \u2192 RAV verifies \u2192 escrow holds funds",
        "After checkout + 5 days \u2192 funds released \u2192 Stripe Connect payout",
    ])

    doc.add_heading("Traveler Flow", level=2)
    add_numbered_list(doc, [
        "Browse/search listings (voice search, text chat, filters)",
        "View property details with fair value scoring",
        "Place bids or propose alternate dates",
        "Checkout via Stripe \u2192 booking confirmed",
        "Track booking in My Bookings, file disputes if needed",
    ])

    doc.add_heading("Admin Flow", level=2)
    add_numbered_list(doc, [
        "Dashboard with tabs: Users, Listings, Bookings, Escrow, Payouts, Financials, Disputes, Voice",
        "Approve/reject listings and users (now with bulk actions)",
        "Manage escrow lifecycle (verify, hold, release, refund)",
        "Dispute resolution queue with assignment",
        "Voice search monitoring and quota management",
    ])

    add_horizontal_rule(doc)

    # ── Features Built ──
    doc.add_heading("Features Built Across 24+ Sessions", level=1)

    features_data = [
        ("Auth", "Email/password + Google OAuth, role-based access (6 roles), email verification, user approval workflow"),
        ("Listings", "Create/edit listings, nightly pricing, fair value scoring, photo uploads, per-night rate with auto price calculation"),
        ("Bidding", "Bid on listings, propose alternate dates, 24hr expiry, owner accept/reject/counter"),
        ("Booking", "Stripe Checkout, fee breakdown (base + service + cleaning + tax), booking confirmation flow"),
        ("Payments", "Stripe Connect (owner onboarding + payouts), webhooks (6 events), escrow management"),
        ("Cancellation", "Policy-based (flexible/moderate/strict/super_strict) renter cancellation, owner cancellation with full refund, Stripe refunds"),
        ("Escrow", "6-status lifecycle, owner confirmation, RAV verification, auto-release after checkout+5d, hold/unhold, refund"),
        ("Disputes", "Renter can file disputes, admin queue with assignment, resolution with refund"),
        ("Voice Search", "VAPI integration, tier-based quotas, admin overrides, usage dashboard, search logging"),
        ("Text Chat", "RAVIO AI assistant via OpenRouter"),
        ("Calculator", "Maintenance fee breakeven calculator for 9 brands"),
        ("Travel Requests", "Travelers post what they want, auto-matched when listings appear"),
        ("Owner Dashboard", "Earnings, bookings, listings management, Stripe Connect status, escrow visibility"),
        ("Admin Dashboard", "8-tab dashboard with cross-entity linking, search, date filters, bulk actions, notes, age badges, dispute assignment"),
        ("Executive Dashboard", "Marketplace health metrics, industry feed"),
        ("SEO", "Meta tags, sitemap, robots.txt, FAQ JSON-LD, OG images"),
        ("Security", "CSP headers, rate limiting (7 edge functions), RLS policies"),
        ("GDPR", "Data export, account deletion with 14-day grace period, cookie consent"),
        ("Architecture", "Auto-generated flow diagrams from declarative manifests"),
        ("PWA", "Service worker, installable, offline-capable"),
    ]

    add_table_from_data(
        doc,
        ["Area", "What\u2019s Built"],
        features_data,
    )

    add_horizontal_rule(doc)

    # ── Current Numbers ──
    doc.add_heading("Current Numbers", level=1)

    add_table_from_data(
        doc,
        ["Metric", "Count"],
        [
            ("Automated tests", "402 (all passing)"),
            ("Database migrations", "31 (DEV), 23 (PROD)"),
            ("Edge functions", "24"),
            ("Supabase RLS policies", "Extensive across all tables"),
            ("Pages / routes", "~20"),
            ("Commits on dev ahead of main", "Many \u2014 needs a PR to merge"),
        ],
    )

    add_horizontal_rule(doc)

    # ── Remaining Pre-Launch Items ──
    doc.add_heading("Remaining Pre-Launch Items", level=1)
    add_body(doc, "6 open issues remain before the platform can go live:")

    add_table_from_data(
        doc,
        ["#", "Issue", "Status"],
        [
            ("#127", "Business Formation & Stripe Tax Activation", "Blocked on LLC / EIN"),
            ("#87", "Launch readiness checklist", "Ready when other items close"),
            ("#80", "Legal review: ToS and Privacy Policy", "Needs lawyer review"),
            ("#74", "Google Analytics (GA4) Integration", "Not started"),
            ("#64", "1099-K Compliance", "Not started"),
            ("#62", "Admin Tax Reporting", "Not started"),
        ],
    )

    add_horizontal_rule(doc)

    # ── Current Platform State ──
    doc.add_heading("Current Platform State", level=1)
    add_bullet_list(doc, [
        "PROD: Staff Only Mode enabled \u2014 platform locked for internal testing",
        "Stripe Tax: Code ready but not activated in Stripe Dashboard (blocked on #127)",
        "GitHub Actions: Issue Notifications workflow temporarily disabled (Resend quota)",
        "Supabase CLI: Currently linked to DEV project",
    ])

    # Footer
    add_footer(doc, "Rent-A-Vacation \u2022 Confidential \u2022 Generated " + datetime.now().strftime("%B %d, %Y"))

    # Save
    output_path = os.path.join(SCRIPT_DIR, f"RAV-Platform-Overview-{today}.docx")
    doc.save(output_path)
    print(f"Generated: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_platform_overview()
