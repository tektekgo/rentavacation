"""
Generate RAV Technology Inventory .docx using brand styling from generate_docx.py.
"""

import os
import sys
from datetime import datetime

# Add parent to path so we can import from generate_docx
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from generate_docx import (
    create_branded_doc, add_logo_header, add_page_numbers,
    add_metadata, add_body, add_horizontal_rule, add_table_from_data,
    add_footer, add_blockquote,
    DEEP_TEAL, WARM_CORAL, DARK_NAVY, WHITE, BRAND_FONT,
)
from docx.shared import Pt, RGBColor


SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(SCRIPT_DIR, "RAV-Technology-Inventory-02262026.docx")


def generate():
    doc = create_branded_doc("Technology Inventory")
    add_logo_header(doc, doc_title="Technology & Tools Inventory")
    add_page_numbers(doc)

    add_metadata(doc, [
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Version", "1.0"),
        ("Classification", "Internal \u2014 Business Operations"),
        ("Prepared by", "RAV Engineering Team"),
    ])

    # ── SECTION 1: SUBSCRIPTION / PAID SERVICES ──
    doc.add_heading("1. Subscription & Paid Services", level=1)
    add_body(doc,
        "These are the external SaaS products and services that RAV depends on. "
        "Most operate on free tiers during pre-launch; costs scale with usage post-launch."
    )

    add_table_from_data(doc,
        ["#", "Service", "Purpose", "Free Tier", "Paid Tier", "Status"],
        [
            ["1", "Claude Max (Anthropic)", "AI coding assistant (Claude Code)", "\u2014", "$100\u2013200/mo", "Active"],
            ["2", "Supabase", "Database, Auth, Edge Functions, Storage", "500MB DB, 50K users, 500K invocations", "$25/mo (Pro)", "Active (DEV + PROD)"],
            ["3", "Vercel", "Frontend hosting, CDN, preview deploys", "100GB bandwidth, 6K build min/mo", "$20/mo (Pro)", "Active"],
            ["4", "Stripe", "Payments, Connect payouts, Tax", "No platform fee", "2.9% + $0.30 per txn", "Active"],
            ["5", "VAPI.ai", "Voice assistant (Deepgram + ElevenLabs + OpenAI)", "10 min/mo", "~$0.05\u20130.15/min", "Active"],
            ["6", "OpenRouter", "Text chat LLM (Gemini 3 Flash)", "Free tier models", "~$0.50/M tokens", "Active"],
            ["7", "Resend", "Transactional email", "3,000 emails/mo, 1 domain", "$20/mo (5K+)", "Active"],
            ["8", "Sentry", "Error monitoring & performance", "5K errors/mo, 10K transactions", "$26/mo (Team)", "Active"],
            ["9", "Cloudflare", "DNS, email routing (catch-all)", "Free plan", "\u2014", "Active (free)"],
            ["10", "Percy (BrowserStack)", "Visual regression testing", "5K screenshots/mo", "$399/mo (Team)", "Active (CI)"],
            ["11", "Qase.io", "Test case management & reporting", "500 test cases, 3 users", "$36/mo (Startup)", "Active (CI)"],
            ["12", "NewsAPI", "Industry news feed (exec dashboard)", "100 req/day (dev only)", "$449/mo (Business)", "Active (dev key)"],
            ["13", "PostHog", "Product analytics & user behavior", "1M events/mo", "$1,600/yr (Scale Add-on)", "Active (coupon thru Jan 2027)"],
            ["14", "AirDNA", "Vacation rental market intelligence (BYOK)", "\u2014", "~$250\u2013500/mo (user-paid)", "Planned \u2014 BYOK model"],
            ["15", "STR Global", "Short-term rental benchmarks (BYOK)", "\u2014", "Custom pricing (user-paid)", "Planned \u2014 BYOK model"],
            ["16", "Canva", "Marketing assets, brand design", "Free tier", "$13/mo (Pro)", "Used for design"],
            ["17", "GitHub", "Repo, Issues, Actions CI/CD", "Free (public), 2K Actions min/mo", "$4/user/mo (Team)", "Active"],
        ],
    )

    # ── SECTION 2: INTEGRATED PROVIDERS ──
    doc.add_heading("2. Integrated Providers", level=1)
    add_body(doc,
        "These are billed through a primary service listed above \u2014 you do not pay them separately."
    )

    add_table_from_data(doc,
        ["Provider", "Billed Via", "What It Does"],
        [
            ["Deepgram", "VAPI", "Speech-to-text (STT)"],
            ["ElevenLabs", "VAPI", "Text-to-speech (TTS)"],
            ["OpenAI GPT-4o-mini", "VAPI", "Voice assistant LLM"],
            ["Google Gemini 3 Flash", "OpenRouter", "Text chat LLM"],
            ["Google OAuth 2.0", "Supabase Auth", "Social login (Sign in with Google)"],
            ["PostgreSQL", "Supabase", "Relational database engine"],
            ["Deno Deploy", "Supabase", "Edge function runtime"],
            ["Let\u2019s Encrypt", "Vercel", "SSL/TLS certificates"],
        ],
    )

    # ── SECTION 3: FREE EXTERNAL APIS ──
    doc.add_heading("3. Free External APIs", level=1)
    add_body(doc,
        "No account or subscription required for these services."
    )

    add_table_from_data(doc,
        ["Service", "What It Does", "Notes"],
        [
            ["FRED API (Federal Reserve)", "Economic indicators for exec dashboard", "Public API, no key required"],
            ["Google Fonts", "Roboto font family", "CDN-hosted, free"],
            ["Unsplash", "Stock property & destination photos", "Free for commercial use with attribution"],
        ],
    )

    # ── SECTION 4: NOT YET INTEGRATED ──
    doc.add_heading("4. Not Yet Integrated in Code", level=1)
    add_body(doc,
        "These services have accounts, code scaffolding, or open GitHub issues but are not yet fully wired into the application."
    )

    add_table_from_data(doc,
        ["Service", "Purpose", "Integration Status", "GitHub Issue"],
        [
            ["Google Analytics (GA4)", "Traffic analytics & marketing attribution", "Not implemented", "#74"],
            ["AirDNA", "Live market data for exec dashboard", "BYOK: edge function + settings UI exist, awaiting user API key", "\u2014"],
            ["STR Global", "Live rental benchmarks for exec dashboard", "BYOK: edge function + settings UI exist, awaiting user API key", "\u2014"],
        ],
    )

    add_body(doc, "")
    add_blockquote(doc,
        "BYOK (Bring Your Own Key): AirDNA and STR Global use a model where the admin enters their own "
        "API key via the Executive Dashboard \u2192 Integration Settings panel. RAV does not pay for these "
        "subscriptions \u2014 they are user-paid. The platform shows demo data until a key is provided."
    )

    # ── SECTION 5: OPEN SOURCE / DEV TOOLS ──
    doc.add_heading("5. Open Source & Development Tools", level=1)
    add_body(doc,
        "Free, open-source tools used in the development stack. No subscription cost."
    )

    add_table_from_data(doc,
        ["Category", "Tools"],
        [
            ["Frontend Framework", "React 18, TypeScript 5.8, Vite 5.4"],
            ["Styling", "Tailwind CSS 3.4, shadcn/ui (Radix UI primitives)"],
            ["State & Data", "TanStack React Query, React Hook Form, Zod"],
            ["UI Components", "Lucide icons, Recharts, Mermaid, date-fns, Embla Carousel"],
            ["Testing", "Vitest, Playwright, Testing Library, jsdom"],
            ["Code Quality", "ESLint 9, Husky, lint-staged"],
            ["PWA", "vite-plugin-pwa, Workbox (offline support, installable app)"],
        ],
    )

    # ── SECTION 6: MONTHLY COST ESTIMATE ──
    doc.add_heading("6. Monthly Cost Estimate (Pre-Launch)", level=1)
    add_body(doc,
        "Estimated costs while the platform is in pre-launch / Staff Only Mode with minimal traffic."
    )

    add_table_from_data(doc,
        ["Service", "Estimated Cost", "Notes"],
        [
            ["Claude Max", "$100\u2013200", "Primary development tool"],
            ["Supabase (2 projects)", "$0\u201350", "Free tier covers pre-launch"],
            ["Vercel", "$0\u201320", "Free tier likely sufficient initially"],
            ["Stripe", "$0", "Only charges per transaction"],
            ["VAPI", "$0\u201310", "Minimal voice usage pre-launch"],
            ["OpenRouter", "$0\u20135", "Gemini Flash is very cheap"],
            ["Resend", "$0", "Free tier (3,000 emails/mo)"],
            ["Sentry", "$0", "Free tier (5K errors/mo)"],
            ["Cloudflare", "$0", "Free plan"],
            ["Percy", "$0", "Free tier for CI"],
            ["Qase", "$0", "Free tier"],
            ["PostHog", "$0", "Coupon until Jan 2027; then ~$133/mo ($1,600/yr)"],
            ["AirDNA", "$0", "BYOK \u2014 user-paid, not a RAV expense"],
            ["STR Global", "$0", "BYOK \u2014 user-paid, not a RAV expense"],
            ["NewsAPI", "$0", "Dev key (prod needs $449/mo or alternative)"],
            ["GitHub", "$0", "Free for current usage"],
            ["Canva", "$0\u201313", "Optional"],
            ["TOTAL", "$100\u2013300/mo", "Pre-launch; mostly Claude Max"],
        ],
    )

    add_body(doc, "")
    add_blockquote(doc,
        "Post-launch costs scale with usage \u2014 mainly Stripe (per-transaction), "
        "Supabase (DB size + edge invocations), Vercel (bandwidth), and VAPI (voice minutes). "
        "PostHog coupon expires Jan 21, 2027 \u2014 then $1,600/yr. "
        "AirDNA and STR Global are BYOK (user-paid, not a RAV expense)."
    )

    # ── SECTION 7: ENVIRONMENT CONFIGURATION ──
    doc.add_heading("7. Environment Configuration Summary", level=1)

    doc.add_heading("Frontend Environment Variables (.env.local / Vercel)", level=2)
    add_table_from_data(doc,
        ["Variable", "Service", "Set In"],
        [
            ["VITE_SUPABASE_URL", "Supabase", ".env.local + Vercel"],
            ["VITE_SUPABASE_ANON_KEY", "Supabase", ".env.local + Vercel"],
            ["VITE_VAPI_PUBLIC_KEY", "VAPI", ".env.local + Vercel"],
            ["VITE_VAPI_ASSISTANT_ID", "VAPI", ".env.local + Vercel"],
            ["VITE_SENTRY_DSN", "Sentry", ".env.local + Vercel"],
            ["VITE_FEATURE_VOICE_ENABLED", "Feature flag", ".env.local + Vercel"],
        ],
    )

    doc.add_heading("Backend Secrets (Supabase Edge Functions)", level=2)
    add_table_from_data(doc,
        ["Secret", "Service", "Used By"],
        [
            ["STRIPE_SECRET_KEY", "Stripe", "Checkout, payouts, refunds, webhooks"],
            ["STRIPE_WEBHOOK_SECRET", "Stripe", "Webhook signature verification"],
            ["RESEND_API_KEY", "Resend", "All transactional emails"],
            ["OPENROUTER_API_KEY", "OpenRouter", "Text chat (RAVIO)"],
            ["NEWSAPI_KEY", "NewsAPI", "Industry news feed (optional)"],
            ["SUPABASE_SERVICE_ROLE_KEY", "Supabase", "Admin operations in edge functions"],
            ["IS_DEV_ENVIRONMENT", "Internal", "Guards seed data in production"],
        ],
    )

    doc.add_heading("CI/CD Secrets (GitHub Actions)", level=2)
    add_table_from_data(doc,
        ["Secret", "Service", "Used By"],
        [
            ["PERCY_TOKEN", "Percy", "Visual regression tests"],
            ["QASE_API_TOKEN", "Qase", "Test reporting"],
            ["SUPABASE_URL", "Supabase", "CI test environment"],
            ["SUPABASE_ANON_KEY", "Supabase", "CI test environment"],
            ["RESEND_GITHUB_NOTIFICATIONS_KEY", "Resend", "Issue email notifications"],
        ],
    )

    # ── SECTION 8: ARCHITECTURE OVERVIEW ──
    doc.add_heading("8. Architecture Overview", level=1)

    add_table_from_data(doc,
        ["Layer", "Technology", "Details"],
        [
            ["Frontend", "React + TypeScript + Vite", "SPA deployed on Vercel CDN"],
            ["UI Library", "Tailwind CSS + shadcn/ui", "Utility-first CSS + Radix primitives"],
            ["Backend", "Supabase", "PostgreSQL + PostgREST API + Auth + Edge Functions"],
            ["Edge Functions", "Deno (TypeScript)", "24 serverless functions for business logic"],
            ["Payments", "Stripe", "Checkout, Connect (owner payouts), webhooks, tax"],
            ["Voice AI", "VAPI \u2192 Deepgram + ElevenLabs + OpenAI", "Browser-based voice search"],
            ["Text AI", "OpenRouter \u2192 Gemini 3 Flash", "Conversational assistant (RAVIO)"],
            ["Email", "Resend", "Transactional emails (7 templates)"],
            ["Monitoring", "Sentry", "Error tracking + performance"],
            ["Analytics", "PostHog", "Product analytics + user behavior (Scale plan)"],
            ["Market Data", "AirDNA + STR Global", "Rental market intelligence (planned)"],
            ["DNS / CDN", "Cloudflare + Vercel", "DNS routing, email catch-all, edge CDN"],
            ["CI/CD", "GitHub Actions", "Lint, test, visual regression, deploy"],
            ["Testing", "Vitest + Playwright + Percy + Qase", "Unit, E2E, visual, reporting"],
        ],
    )

    # ── FOOTER ──
    add_footer(doc, f"Rent-A-Vacation \u2022 Technology Inventory \u2022 Confidential \u2022 {datetime.now().strftime('%B %Y')}")

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    generate()
