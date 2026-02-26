"""
Generate branded .docx files for RAV roadmap and status report.
Uses brand colors from BRAND-STYLE-GUIDE.md:
  - Deep Teal: #1C7268
  - Warm Coral: #E8703A
  - Background: #F8F6F3
  - Dark Navy: #1D2E38
  - Font: Roboto (falls back to Calibri on systems without Roboto)
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Brand colors
DEEP_TEAL = RGBColor(0x1C, 0x72, 0x68)
WARM_CORAL = RGBColor(0xE8, 0x70, 0x3A)
DARK_NAVY = RGBColor(0x1D, 0x2E, 0x38)
LIGHT_BG = RGBColor(0xF8, 0xF6, 0xF3)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1C7268"
TABLE_ALT_ROW = "F0F7F6"
TABLE_BORDER = "CCCCCC"
BRAND_FONT = "Calibri"  # Roboto not always installed; Calibri is professional and universal

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, "..", ".."))
RAVIO_LOGO = os.path.join(PROJECT_ROOT, "public", "ravio-the-chat-genie-128px.png")


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcBorders_existing = tcPr.find(qn('w:tcBorders'))
    if tcBorders_existing is not None:
        tcPr.remove(tcBorders_existing)
    tcPr.append(tcBorders)


def style_table(table, has_header=True):
    """Style a table with brand colors."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = BRAND_FONT
                    run.font.size = Pt(9)
            if i == 0 and has_header:
                set_cell_shading(cell, TABLE_HEADER_BG)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = WHITE
                        run.font.bold = True
            elif i % 2 == 0 and i > 0:
                set_cell_shading(cell, TABLE_ALT_ROW)


def create_branded_doc(title):
    """Create a new document with brand styling."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = BRAND_FONT
    font.size = Pt(10)
    font.color.rgb = DARK_NAVY

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style headings
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = BRAND_FONT
        heading_style.font.color.rgb = DEEP_TEAL
        if level == 1:
            heading_style.font.size = Pt(20)
            heading_style.font.bold = True
        elif level == 2:
            heading_style.font.size = Pt(15)
            heading_style.font.bold = True
        elif level == 3:
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True

    return doc


def add_page_numbers(doc):
    """Add page numbers to the document footer."""
    from datetime import datetime
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # "Page X of Y" using Word field codes
        run1 = p.add_run("Page ")
        run1.font.name = BRAND_FONT
        run1.font.size = Pt(8)
        run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # PAGE field
        fld_begin = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>')
        p._p.append(fld_begin)
        fld_code = parse_xml(f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
        p._p.append(fld_code)
        fld_end = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>')
        p._p.append(fld_end)

        run2 = p.add_run(" of ")
        run2.font.name = BRAND_FONT
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # NUMPAGES field
        fld_begin2 = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>')
        p._p.append(fld_begin2)
        fld_code2 = parse_xml(f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>')
        p._p.append(fld_code2)
        fld_end2 = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>')
        p._p.append(fld_end2)


def add_logo_header(doc, doc_title=None):
    """Add branded logo header with RAVIO."""
    # RAV brand name as styled text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("RENT-A-VACATION")
    run.font.name = BRAND_FONT
    run.font.size = Pt(28)
    run.font.color.rgb = DEEP_TEAL
    run.font.bold = True

    # Document title
    if doc_title:
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pt.paragraph_format.space_before = Pt(2)
        pt.paragraph_format.space_after = Pt(2)
        run_t = pt.add_run(doc_title)
        run_t.font.name = BRAND_FONT
        run_t.font.size = Pt(16)
        run_t.font.color.rgb = DARK_NAVY
        run_t.font.bold = True

    # Tagline
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run("Name Your Price. Book Your Paradise.")
    run2.font.name = BRAND_FONT
    run2.font.size = Pt(11)
    run2.font.color.rgb = WARM_CORAL
    run2.font.italic = True

    # RAVIO chatbot logo + text
    if os.path.exists(RAVIO_LOGO):
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(8)
        run3 = p3.add_run()
        run3.add_picture(RAVIO_LOGO, width=Inches(0.4))
        run4 = p3.add_run("  Ask RAVIO")
        run4.font.name = BRAND_FONT
        run4.font.size = Pt(11)
        run4.font.color.rgb = DEEP_TEAL
        run4.font.bold = True
        run5 = p3.add_run("  \u2014  Just Say Where. RAVIO Does the Rest.")
        run5.font.name = BRAND_FONT
        run5.font.size = Pt(9)
        run5.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run5.font.italic = True

    # Horizontal rule
    add_horizontal_rule(doc)


def add_horizontal_rule(doc):
    """Add a teal horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="1C7268"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_metadata(doc, pairs):
    """Add key-value metadata lines."""
    for key, value in pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run_key = p.add_run(f"{key}: ")
        run_key.font.name = BRAND_FONT
        run_key.font.size = Pt(10)
        run_key.font.bold = True
        run_key.font.color.rgb = DARK_NAVY
        run_val = p.add_run(value)
        run_val.font.name = BRAND_FONT
        run_val.font.size = Pt(10)
        run_val.font.color.rgb = DARK_NAVY


def add_body(doc, text, bold=False, italic=False, size=10, color=None):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BRAND_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or DARK_NAVY
    return p


def add_blockquote(doc, text):
    """Add a styled blockquote."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="1C7268"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = BRAND_FONT
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table_from_data(doc, headers, rows):
    """Add a styled table from header list and row list."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # Headers
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
    style_table(table)
    return table


def add_footer(doc, text):
    """Add a footer paragraph."""
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = BRAND_FONT
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ============================================================
# ROADMAP DOCUMENT
# ============================================================

def generate_roadmap():
    doc = create_branded_doc("Roadmap")
    add_logo_header(doc, doc_title="Product Roadmap & Technical Overview \u2014 Draft")
    add_page_numbers(doc)

    # Metadata
    add_metadata(doc, [
        ("Date", "February 22, 2026"),
        ("Version", "v0.9.0"),
        ("Status", "Pre-Launch (Staff Only Mode \u2014 all features deployed, platform locked for internal testing)"),
        ("Last Updated", "February 22, 2026 at 11:30 PM EST"),
    ])
    doc.add_paragraph()

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_heading("The Problem", level=2)
    add_body(doc, "The vacation ownership (timeshare) industry is valued at $10.5 billion, yet owners of unused weeks have no efficient, trusted marketplace to monetize them. Existing options are fragmented \u2014 classified ads, Facebook groups, or legacy resale sites \u2014 with no pricing transparency, no buyer protection, and no tools for owners to manage their inventory. Meanwhile, travelers looking for vacation club properties have no way to discover available weeks, negotiate pricing, or book with confidence.")
    doc.add_heading("The Solution", level=2)
    add_body(doc, "Rent-A-Vacation (RAV) is a peer-to-peer vacation rental marketplace purpose-built for vacation club and timeshare owners. The platform creates a two-sided marketplace where owners list unused timeshare weeks and travelers discover, negotiate, and book vacation rentals \u2014 with transparent per-night pricing, a bidding engine that lets travelers propose their own terms, and trust infrastructure that protects both sides of every transaction.")
    add_body(doc, "The platform is feature-complete for MVP across 19 completed development phases, with 306 automated tests passing, zero type errors, and zero lint errors. All 21 database migrations and 17 edge functions are deployed to both development and production environments.")
    doc.add_paragraph()
    add_body(doc, "Key Differentiators:", bold=True)
    bullets = [
        ("Two-Sided Marketplace with Real-Time Negotiation:", " Travelers can book at listed prices, bid their own price (\"Name Your Price\"), propose different dates, or post wish lists (\"Vacation Wishes\") that owners compete to fulfill. Owners see live demand signals while creating listings. Auto-matching connects newly approved listings with open traveler requests."),
        ("Traveler-Friendly Pricing:", " Per-night rate transparency (not lump-sum), flexible date proposals that auto-compute from nightly rate, and AI-powered fair value analysis so travelers know if a price is competitive."),
        ("Owner-Centric Tools:", " Full business intelligence suite (\"Owner's Edge\") with earnings tracking against maintenance fee targets, pricing recommendations based on comparable accepted bids, bid activity feed, and idle week alerts."),
        ("Trust & Payment Protection:", " Escrow system (PaySafe) holds funds until the traveler physically checks in. Owner verification (TrustShield) with progressive trust levels. Admin-controlled approval workflows. 4 cancellation policy tiers."),
        ("AI-Enhanced Search:", " Voice concierge (Ask RAVIO) and text chat (Chat with RAVIO) provide natural language property search as an additional discovery channel \u2014 complementing the traditional search, filter, and browse experience."),
    ]
    for label, desc in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(label)
        run_b.font.bold = True
        run_b.font.name = BRAND_FONT
        run_b.font.size = Pt(10)
        run_d = p.add_run(desc)
        run_d.font.name = BRAND_FONT
        run_d.font.size = Pt(10)

    # 2. Platform Capabilities
    doc.add_heading("2. Platform Capabilities — BUILT", level=1)

    # 2.1 Core Marketplace
    doc.add_heading("2.1 Core Marketplace", level=2)
    add_table_from_data(doc,
        ["Capability", "Description", "Status"],
        [
            ["User Authentication", "Email/password + Google OAuth with admin approval workflow", "BUILT"],
            ["Role-Based Access", "5 roles: RAV Owner, RAV Admin, RAV Staff, Property Owner, Renter", "BUILT"],
            ["Property Registration", "Multi-step form with resort search (117 resorts, 351 unit types), auto-populate specs, image upload", "BUILT"],
            ["Listing Management", "Draft → Pending → Admin Approval → Active lifecycle, per-night pricing, cancellation policies", "BUILT"],
            ["Booking Flow", "Browse → View → Book Now → Stripe Checkout → Payment Capture → Booking Confirmation", "BUILT"],
            ["Escrow System (PaySafe)", "Funds held until check-in confirmed, released to owner after checkout + 5 days", "BUILT"],
            ["Owner Confirmation Timer", "Configurable countdown (default 60 min), up to 2 extensions of 30 min, auto-cancel on timeout", "BUILT"],
            ["Cancellation Policies", "4 tiers: Flexible (100% ≥1 day), Moderate (100% ≥5 days), Strict (50% ≥7 days), Super Strict (no refunds)", "BUILT"],
            ["Check-in Confirmation", "Traveler confirms arrival or reports issues; issue resolution workflow for admin", "BUILT"],
        ]
    )

    # 2.2 Bidding & Negotiation
    doc.add_heading("2.2 Bidding & Negotiation", level=2)
    add_table_from_data(doc,
        ["Feature", "Description", "Status"],
        [
            ["Place a Bid (Name Your Price)", "Travelers bid on any listing where the owner has opted in. Owners review, accept, reject, or counter-offer", "BUILT"],
            ["Date Proposals", "Travelers propose different dates; bid amount auto-computes from nightly rate \u00d7 proposed nights. Owners see proposed dates highlighted in bid manager", "BUILT"],
            ["Travel Requests (Vacation Wishes)", "Reverse auction \u2014 travelers post dream trips (destination, dates, budget, bedrooms), owners compete with proposals", "BUILT"],
            ["Inspired Requests", "\"Request Similar Dates\" from any listing detail page \u2014 pre-fills destination, dates, bedrooms. Optional \"Send to this owner first\" targeting", "BUILT"],
            ["Auto-Matching", "Newly approved listings are automatically matched against open travel requests by destination, dates (\u00b130 days), budget, bedrooms, and brand", "BUILT"],
            ["Demand Signals", "Owners see matching travel request count + max disclosed budget while creating listings, helping them price competitively", "BUILT"],
        ]
    )

    # 2.3 Pricing & Revenue
    doc.add_heading("2.3 Pricing & Revenue", level=2)
    add_table_from_data(doc,
        ["Aspect", "Detail"],
        [
            ["Atomic Pricing Unit", "Per-night rate (nightly_rate) — all prices computed from this base"],
            ["Platform Commission", "15% default (admin-configurable via System Settings)"],
            ["Pro Owner Discount", "13% commission (−2%)"],
            ["Business Owner Discount", "10% commission (−5%)"],
            ["Stripe Processing", "~2.9% absorbed by RAV within service fee margin"],
            ["Payout Timing", "Owner receives payout after checkout date + 5 days"],
        ]
    )

    # 2.4 Business Intelligence
    doc.add_heading("2.4 Business Intelligence", level=2)
    add_table_from_data(doc,
        ["Dashboard", "Audience", "Description"],
        [
            ["Owner Dashboard (Owner's Edge)", "Property Owners", "6 BI sections: Headline Stats (earned YTD, fees covered %, active bids), Earnings Timeline (chart with maintenance fee target line), My Listings Table (status badges, idle week alerts), Bid Activity Feed, Pricing Intelligence (per-listing fair value + market range), Maintenance Fee Tracker (coverage progress bar)"],
            ["Fair Value Score (RAV SmartPrice)", "All users", "Analysis of comparable accepted bids using P25-P75 percentile range. Shows whether a listing is priced below market, at fair value, or above market. Different messaging for owners vs travelers"],
            ["Maintenance Fee Calculator", "Public (no auth)", "Break-even analysis for 9 vacation club brands and 4 unit types \u2014 shows owners how many weeks to rent to cover annual maintenance fees. Live progress bars, CTA to owner signup"],
            ["Executive Dashboard (RAV Command)", "RAV Leadership", "Investor-grade strategic dashboard. 6 sections: KPI headline bar, Business Performance (4 charts), Marketplace Health (proprietary Liquidity Score and Bid Spread Index), Market Intelligence (AirDNA + STR via BYOK), Industry Feed, Unit Economics"],
        ]
    )

    # 2.5 AI-Enhanced Search
    doc.add_heading("2.5 AI-Enhanced Search", level=2)
    add_table_from_data(doc,
        ["Feature", "Description", "Status"],
        [
            ["Voice Search (Ask RAVIO)", "Voice concierge powered by VAPI + Deepgram Nova-3. Natural language queries, 300ms endpointing, smart denoising. Tier-based daily limits with admin overrides", "BUILT"],
            ["Text Chat (Chat with RAVIO)", "LLM-powered text assistant (OpenRouter / Gemini 3 Flash) with SSE streaming and tool calling. Context-aware across 4 page types", "BUILT"],
            ["Resort Knowledge Base (ResortIQ)", "Database of 117 partner resorts and 351 unit types from 9 vacation club brands. Auto-populates listing specs when owners create listings", "BUILT"],
        ]
    )

    # 2.6 Admin & Operations
    doc.add_heading("2.6 Admin & Operations", level=2)
    add_table_from_data(doc,
        ["Capability", "Description"],
        [
            ["Admin Dashboard", "12 tabs: Overview, Users, Listings (approval workflow), Bookings, Properties, Verifications, Escrow, Payouts, Financials, Issues, Voice, Memberships"],
            ["Voice Admin Controls", "Global config display, tier quota manager, per-user overrides (disable/custom quota), usage dashboard with charts + top users, observability (search log viewer + alert thresholds)"],
            ["Staff Only Mode", "Pre-launch platform lock — 3-layer enforcement (database RLS, Login page, Signup page). Toggle in Admin > System Settings"],
            ["Owner Verification (TrustShield)", "Document upload (deed, certificate, ID), trust levels (new → verified → trusted → premium), admin review workflow"],
            ["Seed Data System", "DEV-only 3-layer system: 8 foundation users (never wiped), 10 properties + 30 listings, 50 renters + 110 bookings + 20 bids. Production-guarded"],
        ]
    )

    # 2.7 Communication
    doc.add_heading("2.7 Communication", level=2)
    add_body(doc, "17 transactional email types via Resend API (notifications@updates.rent-a-vacation.com):", bold=True)
    add_table_from_data(doc,
        ["Category", "Emails"],
        [
            ["Account", "Welcome, User Approved, User Rejected"],
            ["Listings", "Listing Approved, Listing Rejected, Listing Submitted (to admin)"],
            ["Bookings", "Booking Confirmed, Check-in Reminder"],
            ["Owner Confirmation", "Confirmation Request, Extension Notification, Confirmation Timeout"],
            ["Cancellation", "Submitted, Approved, Denied, Counter-Offer"],
            ["Verification", "Document Uploaded (to admin)"],
            ["Support", "Contact Form Submission"],
        ]
    )
    add_body(doc, "In-app notifications with real-time badge count, auto-refresh every 30 seconds.", italic=True, size=9)

    # 3. Membership Tiers
    doc.add_heading("3. Membership Tiers", level=1)
    doc.add_heading("3.1 Renter Tiers", level=2)
    add_table_from_data(doc,
        ["Tier", "Monthly Price", "Voice Searches/Day", "Key Benefits"],
        [
            ["Free", "$0", "5", "Browse listings, place bids, post travel requests"],
            ["Plus", "$9.99", "25", "Priority support, saved searches"],
            ["Premium", "$24.99", "Unlimited", "Early access to new listings, concierge service"],
        ]
    )
    doc.add_heading("3.2 Owner Tiers", level=2)
    add_table_from_data(doc,
        ["Tier", "Monthly Price", "Commission Rate", "Key Benefits"],
        [
            ["Free", "$0", "15% (default)", "List properties, basic dashboard, bid management"],
            ["Pro", "$19.99", "13% (−2% discount)", "Analytics, priority listing placement"],
            ["Business", "$49.99", "10% (−5% discount)", "Multi-property management, API access, dedicated support"],
        ]
    )
    add_blockquote(doc, "Source: Migration 011 (membership_tiers table). Commission rate is admin-configurable in System Settings.")

    # 4. Supported Brands
    doc.add_heading("4. Supported Vacation Club Brands (9)", level=1)
    add_table_from_data(doc,
        ["#", "Brand", "Resort Coverage"],
        [
            ["1", "Hilton Grand Vacations", "62 resorts"],
            ["2", "Marriott Vacation Club", "40 resorts"],
            ["3", "Disney Vacation Club", "15 resorts"],
            ["4", "Wyndham Destinations", "—"],
            ["5", "Hyatt Residence Club", "—"],
            ["6", "Bluegreen Vacations", "—"],
            ["7", "Holiday Inn Club Vacations", "—"],
            ["8", "WorldMark by Wyndham", "—"],
            ["9", "Other / Independent Resort", "—"],
        ]
    )
    add_blockquote(doc, "Source: VACATION_CLUB_BRANDS constant in calculatorLogic.ts and vacation_club_brand database enum. Total: 117 resorts, 351 unit types across 10+ countries.")

    # 5. Technical Infrastructure
    doc.add_heading("5. Technical Infrastructure", level=1)
    doc.add_heading("5.1 Technology Stack", level=2)
    add_table_from_data(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React 18 + TypeScript + Vite + SWC", "Single-page application with strict typing"],
            ["Styling", "Tailwind CSS + shadcn/ui (Radix primitives)", "Utility-first CSS with accessible component library"],
            ["Routing", "React Router v6", "Client-side routing with protected routes"],
            ["Data Fetching", "TanStack React Query v5", "Server state management, caching, optimistic updates"],
            ["Forms", "React Hook Form + Zod", "Schema-validated forms"],
            ["Auth", "Supabase Auth", "Email/password, Google OAuth, role-based access"],
            ["Database", "Supabase PostgreSQL", "Row Level Security (RLS), pg_cron, pg_net"],
            ["Backend", "Supabase Edge Functions (Deno)", "17 serverless functions"],
            ["Payments", "Stripe Checkout", "Payment capture, escrow, webhooks"],
            ["Email", "Resend API", "Transactional emails with branded HTML templates"],
            ["Voice AI", "VAPI + Deepgram Nova-3", "Voice transcription and natural language processing"],
            ["Text AI", "OpenRouter (Gemini 3 Flash)", "LLM chat with SSE streaming and tool calling"],
            ["Charts", "Recharts", "Dashboard analytics and data visualization"],
            ["Hosting", "Vercel (frontend) + Supabase (backend)", "Auto-deploy from GitHub"],
            ["CI/CD", "GitHub Actions", "Lint, typecheck, unit tests, E2E, Percy visual regression"],
            ["PWA", "vite-plugin-pwa + Workbox", "Service worker, install prompt, offline detection"],
        ]
    )

    # 5.2 Migrations
    doc.add_heading("5.2 Database Migrations (21 total)", level=2)
    add_table_from_data(doc,
        ["Migration", "Purpose"],
        [
            ["001", "Core schema: profiles, user_roles, properties, listings, bookings, RLS"],
            ["002", "Seed data (optional)"],
            ["003", "Bidding system: listing_bids, travel_requests, travel_proposals, notifications"],
            ["004", "Payout tracking fields on bookings + booking_confirmations"],
            ["005", "Cancellation policies + refund calculation function"],
            ["006", "Owner verification + trust levels + platform guarantee fund"],
            ["007–008", "Voice auth (user approval system, voice usage limits)"],
            ["010", "Role upgrade requests system"],
            ["011", "Membership tiers (6 tiers) + voice toggles + commission config"],
            ["012", "Phase 13: property images, owner confirmation timer, system settings"],
            ["013", "Executive dashboard settings (API key storage)"],
            ["014", "Staff Only Mode (pre-launch platform lock)"],
            ["015", "Seed data foundation flag (is_seed_foundation)"],
            ["016", "Fair Value Score RPC (P25-P75 percentile analysis)"],
            ["017", "Owner Dashboard RPCs + maintenance fee columns"],
            ["018", "Travel request enhancement notification types"],
            ["019", "PostgREST FK fix (10 tables redirected to profiles)"],
            ["020", "Per-night pricing (nightly_rate) + date proposals + inspired requests"],
            ["021", "Voice admin: search logs, user overrides, alert thresholds, 3 RPCs"],
        ]
    )

    # 5.3 Edge Functions
    doc.add_heading("5.3 Edge Functions (17 total)", level=2)
    add_table_from_data(doc,
        ["Function", "Trigger", "Purpose"],
        [
            ["create-booking-checkout", "Client call", "Creates Stripe Checkout session with listing details and tier-aware commission"],
            ["verify-booking-payment", "Stripe webhook", "Validates payment, creates booking + booking_confirmation with owner acceptance timer"],
            ["send-email", "Client call", "Generic email dispatch via Resend API"],
            ["send-approval-email", "Client call", "Approval/rejection emails for listings and users (4 variants)"],
            ["send-booking-confirmation-reminder", "Client/internal", "Owner deadline reminders + acceptance notifications"],
            ["send-cancellation-email", "Internal", "Cancellation status notifications (4 variants)"],
            ["send-contact-form", "Client call", "Contact form submission handler"],
            ["send-verification-notification", "Client call", "Admin notification on document upload"],
            ["process-deadline-reminders", "CRON (every 30 min)", "Scan deadlines, send reminders, process owner timeouts, travel request expiry warnings"],
            ["match-travel-requests", "Internal", "Auto-match approved listings to open travel requests (budget-aware, deduped)"],
            ["voice-search", "VAPI webhook", "Property search via voice — shared search module, state name expansion"],
            ["text-chat", "Client call", "OpenRouter LLM with SSE streaming, tool calling, 4 context modes"],
            ["seed-manager", "Client call", "DEV-only 3-layer seed data system (production-guarded)"],
            ["fetch-industry-news", "Client call", "NewsAPI + Google News RSS for exec dashboard (60-min cache)"],
            ["fetch-macro-indicators", "Client call", "FRED consumer confidence + travel data"],
            ["fetch-airdna-data", "Client call", "AirDNA market comparisons (BYOK)"],
            ["fetch-str-data", "Client call", "STR hospitality benchmarks (BYOK)"],
        ]
    )

    # 5.4 Quality Metrics
    doc.add_heading("5.4 Quality Metrics", level=2)
    add_table_from_data(doc,
        ["Metric", "Value"],
        [
            ["Automated Tests", "306 (all passing)"],
            ["Type Errors", "0 (strict TypeScript)"],
            ["Lint Errors", "0 (ESLint)"],
            ["Build Status", "Clean (Vite production build)"],
            ["CI Pipeline", "5-job: lint+typecheck → unit tests → E2E → visual regression → lighthouse"],
            ["Coverage Thresholds", "25% statements, 25% branches, 30% functions, 25% lines"],
        ]
    )

    # 6. Completed Development Phases
    doc.add_heading("6. Completed Development Phases", level=1)

    phases = [
        ("Phase 1: Voice Search", "Nov 2025",
         "VAPI voice assistant integration with natural language property search on the Rentals page. Real-time voice transcription and conversational query refinement.",
         "34% voice adoption rate, 87% search success rate, NPS +68, +23% conversion vs manual search."),
        ("Phase 2: Resort Master Data", "Feb 12, 2026",
         "Imported 117 resorts (Hilton 62, Marriott 40, Disney 15) with 351 unit types. Searchable listing flow with Command component and auto-populate functionality.",
         "Listing completion time reduced from 22 min to 8 min (−64%). Completion rate increased from 67% to 94% (+27%). Owner satisfaction: 4.7 stars (+0.9)."),
        ("Phase 3 (Partial): Voice Auth & Approval", "Feb 15, 2026",
         "Three-phase rollout: authentication gate (voice disabled for unauthenticated users), admin-controlled user approval system with email notifications, and daily voice quota with real-time usage indicator.",
         "Estimated $27K/month API cost savings (90% reduction). Voice abuse prevention with enforced quotas. Full admin control over beta access."),
        ("Phase 4: UI Fixes & Documentation", "Feb 13–15, 2026",
         "Calendar tabs, pagination, favorites system, forgot-password flow, user guide updates, FAQ, how-it-works, and admin documentation.", None),
        ("Phase 5: Core Business Flows", "Feb 13, 2026",
         "Replaced mock data with real Supabase queries. Built complete booking flow: Browse → View → Book → Stripe Checkout → Payment Capture → Confirmation. Build version system in footer.", None),
        ("Phase 6: Role Upgrade System", "Feb 14, 2026",
         "Self-service role upgrade requests with admin approval. Eliminated dead-end UX flows (non-owners seeing empty dashboards, unauthorized bid attempts). Signup role selection (owner vs renter).", None),
        ("Phase 7: UI Excellence & Social Proof", "Feb 14, 2026",
         "Social proof indicators (favorites count, freshness badges, popularity badges), honest content replacement (removed fabricated stats), visual polish (gradients, hover effects, trust indicators), and \"Similar Properties\" recommendations.", None),
        ("Phase 8: Testing Infrastructure", "Feb 14, 2026",
         "Vitest with v8 coverage, Playwright E2E, Percy visual regression, GitHub Actions CI (5-job pipeline), Husky pre-commit hooks, test helpers and fixtures.", None),
        ("Phase 9: Voice Toggles & Membership Tiers", "Feb 14, 2026",
         "6 membership tiers (3 renter + 3 owner), admin voice feature toggles (master + per-feature), configurable platform commission with tier discounts, tier-aware voice quotas.", None),
        ("Phase 10: Additional Improvements", "Feb 15–16, 2026",
         "Contact form, link audit, role terminology standardization (\"Traveler\" → \"Renter\"), UX feedback improvements (inline success states replacing toasts).", None),
        ("Phase 11: Progressive Web App", "Feb 16, 2026",
         "Full PWA support with Workbox service worker (59 precached entries), install prompt, offline detection, iOS meta tags.", None),
        ("Phase 13: Core Business Flow Completion", "Feb 20, 2026",
         "5 tracks: approval email notifications, owner bidding UI, property image upload with drag-and-drop, payout tracking, and owner confirmation timer with extension system.", None),
        ("Phase 14: Executive Dashboard", "Feb 20, 2026",
         "Investor-grade dark-themed strategic dashboard with 6 sections, 4 edge functions for external data (NewsAPI, FRED, AirDNA BYOK, STR BYOK), 4 data hooks, proprietary metrics (Liquidity Score, Bid Spread Index).", None),
        ("Phase 15: Fair Value Score — RAV SmartPrice", "Feb 21, 2026",
         "PostgreSQL RPC function analyzing comparable accepted bids (P25-P75 percentile range). Frontend components with role-specific messaging. Wired into Rentals cards, PropertyDetail sidebar, and owner listings management.", None),
        ("Phase 16: Maintenance Fee Calculator — Fee Freedom Calculator", "Feb 21, 2026",
         "Public break-even analysis tool at /calculator. Pure calculation logic covering 9 brands and 4 unit types. Color-coded progress bars and CTA to owner signup.", None),
        ("Phase 17: Owner Dashboard — Owner's Edge", "Feb 21, 2026",
         "6 business intelligence sections replacing placeholder Overview tab. 2 new PostgreSQL RPCs, 4 data hooks, 6 analytics components including earnings timeline chart and maintenance fee tracker.", None),
        ("Phase 18: Travel Request Enhancements — Vacation Wishes", "Feb 21, 2026",
         "Auto-match engine on listing approval, demand signal display on listing form, \"Post a Travel Request\" CTA on empty search results, and expiry warning system.", None),
        ("Phase 19: Flexible Date Booking + Per-Night Pricing", "Feb 22, 2026",
         "Switched from lump-sum to per-night pricing. Added \"Propose Different Dates\" bidding mode and \"Request Similar Dates\" inspired travel requests from listing detail. Shared pricing utility replacing 4 duplicated functions.", None),
        ("Voice Tracks C-D: Admin Controls + Observability", "Feb 22, 2026",
         "Voice admin dashboard with 5 sections: config info, tier quota manager, per-user overrides, usage dashboard (charts + top users), observability (search log viewer + alert thresholds). Auto-logging of all voice searches.", None),
        ("Content Accuracy Audit", "Feb 22, 2026",
         "Fixed commission rate (10% → 15%) across 7 code files + 3 tests. Corrected brand list (Westgate → WorldMark). Fixed voice quota display (flat → tier-based). Added 9 missing admin documentation sections. Established Content Accuracy policy.", None),
    ]

    for name, date, desc, impact in phases:
        doc.add_heading(name, level=3)
        p = doc.add_paragraph()
        run_date = p.add_run(f"Completed: {date}")
        run_date.font.name = BRAND_FONT
        run_date.font.size = Pt(9)
        run_date.font.italic = True
        run_date.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        add_body(doc, desc)
        if impact:
            p2 = doc.add_paragraph()
            run_label = p2.add_run("Impact (PROJECTED): ")
            run_label.font.bold = True
            run_label.font.name = BRAND_FONT
            run_label.font.size = Pt(9)
            run_label.font.color.rgb = WARM_CORAL
            run_imp = p2.add_run(impact)
            run_imp.font.name = BRAND_FONT
            run_imp.font.size = Pt(9)
            run_imp.font.color.rgb = DARK_NAVY

    # 7. Upcoming Priorities
    doc.add_heading("7. Upcoming Priorities", level=1)
    add_table_from_data(doc,
        ["#", "Phase", "Est. Time", "Pre-Launch?"],
        [
            ["1", "Phase 20A-C: Accounting, Tax & Fee Framework", "14-20h", "A+B Required"],
            ["2", "SEO Optimization", "8-12h", "Recommended"],
            ["3", "Security Hardening (CSP, rate limiting, error monitoring)", "6-10h", "Recommended"],
            ["4", "Phase 20D-F: QuickBooks, 1099-K, Tax Filing", "20-36h", "Post-launch"],
            ["5", "Phase 3: Voice Everywhere", "3-4 weeks", "No"],
            ["6", "Phase 12: Native Mobile (Capacitor)", "2-3 weeks", "No"],
            ["7", "Phase 21: Partial-Week Booking", "20-30h", "No (future)"],
            ["8", "Phase 6: Advanced Features", "TBD", "No (Q3 2026)"],
        ]
    )

    # 8. Current Roadmap (Detail)
    doc.add_heading("8. Current Roadmap — Planned (Detail)", level=1)

    # 8.1 Phase 20
    doc.add_heading("8.1 Phase 20: Accounting, Tax & Fee Framework", level=2)
    add_body(doc, "Priority: Required before public launch (Phases A + B). As a marketplace facilitator in 43+ US states, RAV must collect and remit occupancy/sales taxes before processing real transactions.", bold=False)
    add_table_from_data(doc,
        ["Sub-Phase", "Scope", "Est. Time", "Timeline"],
        [
            ["A: Fee Breakdown", "Separate fee line items on bookings: nightly rate, service fee, cleaning fee, tax. Price breakdown display on PropertyDetail and Checkout.", "4-6h", "Pre-launch"],
            ["B: Stripe Tax Integration", "Auto-calculated occupancy + sales tax at checkout based on property location. Tax line item stored on booking records.", "6-8h", "Pre-launch"],
            ["C: Admin Tax Reporting", "Tax collected report by jurisdiction and month. Owner payout summary. Platform revenue report (service fees only).", "4-6h", "Pre-launch"],
            ["D: QuickBooks Integration", "Sync Stripe transactions to QuickBooks Online via API. Automated revenue recognition and owner payout reconciliation.", "8-12h", "Post-launch"],
            ["E: 1099-K Compliance", "Track owner earnings (>$600/year threshold). Generate 1099-K forms. Owner tax info collection (W-9).", "4-8h", "Before Jan 2027"],
            ["F: Automated Tax Filing", "Avalara or TaxJar integration for auto-filing per jurisdiction. Quarterly remittance reports.", "8-16h", "When volume justifies"],
        ]
    )
    add_blockquote(doc, "Context: Stripe processing fees (~2.9%) are absorbed by RAV within the 15% service fee margin. The platform commission rate (15% default) is admin-configurable. Pro owners pay 13%, Business owners pay 10%.")

    # 8.2 SEO Optimization
    doc.add_heading("8.2 SEO Optimization", level=2)
    add_body(doc, "Priority: Recommended before public launch for organic discovery.", italic=True, size=9)

    add_body(doc, "Planned work:", bold=True)
    add_table_from_data(doc,
        ["Task", "Description", "Est. Time", "Status"],
        [
            ["Per-page meta tags", "Install react-helmet-async, create SEOHead component, add unique title/description to all 15+ public routes", "3-4h", "Planned"],
            ["Sitemap.xml", "Static sitemap with all public routes, add reference in robots.txt", "30min", "Planned"],
            ["Page-level JSON-LD", "FAQPage on /faq, WebApplication on /calculator, Product/Offer on /property/:id, BreadcrumbList, SearchAction", "2-3h", "Planned"],
            ["Image optimization", "Add loading=\"lazy\" + decoding=\"async\", WebP format with fallbacks, responsive srcset", "2-3h", "Planned"],
            ["Route-based code splitting", "Convert static imports to React.lazy() in App.tsx for smaller initial bundle and faster FCP", "1-2h", "Planned"],
            ["Dynamic og:image", "Property-specific Open Graph images when sharing /property/:id on social media", "1h", "Planned"],
            ["404 noindex", "Add noindex meta to NotFound.tsx", "5min", "Planned"],
        ]
    )

    doc.add_paragraph()
    add_body(doc, "What's already built (SEO):", bold=True)
    add_table_from_data(doc,
        ["Feature", "Status", "Details"],
        [
            ["Homepage meta tags", "BUILT", "Title, description, og:*, twitter:* in index.html"],
            ["Organization JSON-LD", "BUILT", "Schema.org markup with social links in index.html"],
            ["robots.txt", "BUILT", "Permissive — allows Googlebot, Bingbot, social crawlers"],
            ["PWA manifest", "BUILT", "Full manifest in vite.config.ts — name, icons, categories"],
            ["Favicons", "BUILT", "ico + png icons for all platforms"],
            ["Alt text on images", "BUILT", "~95% coverage across 23 images"],
            ["Semantic HTML", "BUILT", "h1, h2, main, section, nav used throughout"],
            ["Clean URLs", "BUILT", "RESTful structure with 301 redirects for legacy routes"],
            ["Calculator page title", "BUILT", "Dynamic document.title on /calculator (only page with per-page SEO)"],
            ["Lighthouse CI", "BUILT", "lighthouserc.json config — currently audits 2 URLs"],
        ]
    )
    add_blockquote(doc, "SEO planning docs exist at docs/features/seo-optimization/ with implementation checklists and code examples.")

    # 8.3 Security Hardening
    doc.add_heading("8.3 Security Hardening", level=2)
    add_body(doc, "Priority: Recommended before public launch.", italic=True, size=9)

    add_body(doc, "Planned work:", bold=True)
    add_table_from_data(doc,
        ["Task", "Description", "Est. Time", "Status"],
        [
            ["Content Security Policy", "Add CSP, X-Content-Type-Options, X-Frame-Options, HSTS headers via vercel.json", "2h", "Planned"],
            ["Rate limiting (payment)", "Add per-IP rate limiting to create-booking-checkout and send-email edge functions", "2-3h", "Planned"],
            ["Tighten CORS", "Checkout endpoint currently uses Allow-Origin: * — restrict to production whitelist", "30min", "Planned"],
            ["Error monitoring", "Sentry integration for frontend + edge function error tracking, source map uploads", "2-3h", "Planned"],
            ["Analytics", "Google Analytics 4 or Plausible for page views, events, conversion tracking", "1-2h", "Planned"],
            ["Cookie consent", "GDPR/CCPA consent banner, preference center, conditional analytics loading", "2-3h", "Planned"],
        ]
    )

    doc.add_paragraph()
    add_body(doc, "What's already built (Security):", bold=True)
    add_table_from_data(doc,
        ["Feature", "Status"],
        [
            ["Voice search rate limiting (30 req/min per IP)", "BUILT"],
            ["Text chat rate limiting (60 req/min per IP)", "BUILT"],
            ["CORS whitelist (voice + text-chat)", "BUILT"],
            ["SSL/HTTPS (Vercel + Supabase)", "BUILT"],
            ["Row Level Security (all database tables)", "BUILT"],
            ["Admin approval workflow (user + listing)", "BUILT"],
            ["Staff Only Mode (3-layer enforcement)", "BUILT"],
            ["Terms of Service + Privacy Policy pages", "BUILT"],
        ]
    )

    # 8.4 Phase 21
    doc.add_heading("8.4 Phase 21: Partial-Week Booking", level=2)
    add_body(doc, "Priority: Backlog — deferred until demand validates the pattern through Phase 19's flexible date negotiation.", italic=True, size=9)
    add_body(doc, "Enables travelers to book a subset of an owner's listed dates (e.g., 6 of 8 days). Requires per-night pricing (Phase 19 complete), listing splits, per-split escrow, and handling of cleaning gaps and minimum stay rules. Estimated 20-30 hours.")
    for b in [
        "Owner \"flexible dates\" flag on listings",
        "Calendar subset selection for travelers",
        "Listing splits (booked portion + remaining days become new listing)",
        "Per-segment escrow, confirmation, and payout",
        "Edge cases: cleaning gaps, minimum stay, resort check-in days",
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)

    # 8.5 Voice & Mobile Roadmap
    doc.add_heading("8.5 Voice & Mobile Roadmap", level=2)
    add_table_from_data(doc,
        ["Phase", "Scope", "Prerequisites", "Est. Time"],
        [
            ["Voice Everywhere (Phase 3)", "Voice-assisted listing creation, booking flows, and bidding negotiations", "Voice Tracks C-D complete", "3-4 weeks"],
            ["Native Mobile — Track A", "Capacitor setup + build pipeline", "PWA validates demand", "~2 days"],
            ["Native Mobile — Track B", "Push notifications, camera access, biometric auth", "Track A", "~1 week"],
            ["Native Mobile — Track C", "App Store publishing (Google Play + Apple App Store)", "Track B", "~1 week"],
            ["Native Mobile — Track D", "CI/CD for mobile builds", "Track C", "~2-3 days"],
        ]
    )

    # 8.6 Phase 6
    doc.add_heading("8.6 Phase 6: Advanced Features (Q3 2026)", level=2)
    for b in ["Saved searches & search alerts", "Advanced filtering (map view, amenity search)", "Owner analytics and performance insights", "Calendar integration (Google Calendar, iCal)"]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)

    # 9. Ideas Backlog
    doc.add_heading("9. Ideas Backlog (Unscheduled)", level=1)
    categories = [
        ("Marketing & Growth:", " Blog/content marketing, email campaigns, referral program, social media integration"),
        ("Platform Enhancements:", " Instant booking, dynamic pricing, multi-property management tools, review/rating system"),
        ("Technical:", " Performance optimization, A/B testing framework, CDN for property images, database read replicas"),
        ("Integrations:", " Google Calendar sync, Stripe Connect for payouts, SMS notifications, social login (Facebook, Apple)"),
    ]
    for label, desc in categories:
        p = doc.add_paragraph()
        run_b = p.add_run(label)
        run_b.font.bold = True
        run_b.font.name = BRAND_FONT
        run_b.font.size = Pt(10)
        run_d = p.add_run(desc)
        run_d.font.name = BRAND_FONT
        run_d.font.size = Pt(10)

    # 10. Key Architectural Decisions
    doc.add_heading("10. Key Architectural Decisions", level=1)
    add_table_from_data(doc,
        ["ID", "Decision", "Rationale", "Status"],
        [
            ["DEC-008", "Membership tier & commission architecture", "6 tiers (3 renter, 3 owner) with tier-aware quotas and commission discounts", "Final"],
            ["DEC-011", "PWA first, then Capacitor native shells", "Validate mobile demand before investing in native apps. Existing React codebase carries over — no rewrite needed", "Approved"],
            ["DEC-014", "Separate route for Executive Dashboard", "Different design language (dark-themed) and audience (RAV Owner) from admin dashboard", "Final"],
            ["DEC-015", "BYOK demo/connected pattern for market data", "Honest to investors (no fake data), shows product capability with real integrations", "Final"],
            ["DEC-018", "Staff Only Mode for pre-launch lock", "Global system settings toggle. 3-layer enforcement (DB + Login + Signup). Flip off in admin to go live — no code deploy needed", "Final"],
            ["DEC-019", "Seed Data System", "3-layer edge function approach with foundation user protection. Idempotent, admin UI for one-click reset, production guard via env variable", "Final"],
            ["DEC-020", "Two-tier AI: VAPI voice + OpenRouter text", "Text chat 10-100x cheaper per interaction, works in all environments. Shared search module avoids duplication", "Final"],
            ["DEC-022", "Pricing & Tax Framework", "Per-night pricing + separated fee line items + Stripe Tax before launch + QuickBooks post-launch. Stripe processing fees (~2.9%) absorbed by RAV within the 15% service fee margin", "Approved"],
            ["DEC-023", "Flexible dates: 3-phase approach", "Bid with dates (reuses bidding) > inspired-by requests > partial-week splits. Start lightweight, validate demand, then build full flexibility", "Approved"],
        ]
    )

    # 11. Launch Readiness Checklist
    doc.add_heading("11. Launch Readiness Checklist", level=1)
    add_table_from_data(doc,
        ["Item", "Status", "Notes"],
        [
            ["Core booking flow (Browse > Search > Book > Pay > Confirm > Check-in)", "Ready", "Full Stripe integration with escrow"],
            ["Voice search (Ask RAVIO)", "Ready", "Auth-gated, tier-based quotas, rate-limited, VAPI + Deepgram Nova-3"],
            ["Text chat (Chat with RAVIO)", "Ready", "Deployed on DEV + PROD, OpenRouter key configured"],
            ["Bidding system (Name Your Price)", "Ready", "Full lifecycle: bid > counter > accept > checkout, date proposals"],
            ["Travel requests (Vacation Wishes)", "Ready", "Auto-matching, demand signals, expiry warnings, inspired requests"],
            ["Owner tools (Owner's Edge)", "Ready", "Dashboard, earnings, pricing intel, fee tracker, bid activity"],
            ["Admin suite", "Ready", "12 tabs: approvals, escrow, payouts, voice admin, executive BI"],
            ["Per-night pricing", "Ready", "Phase 19 complete — nightly_rate as atomic pricing unit"],
            ["Voice admin & observability", "Ready", "Voice Tracks C-D complete — admin controls, logging, alerts"],
            ["Fee breakdown display", "In Progress", "Phase 20A — separate service_fee, cleaning_fee, tax line items"],
            ["Stripe Tax integration", "Planned", "Phase 20B — REQUIRED before real transactions"],
            ["SEO optimization", "Planned", "Per-page meta tags, sitemap, JSON-LD, image optimization"],
            ["Error monitoring (Sentry)", "Planned", "Frontend + edge function error tracking"],
            ["Analytics (GA4)", "Planned", "Page views, events, conversion tracking"],
            ["Security headers (CSP)", "Planned", "Content Security Policy, HSTS, X-Frame-Options"],
            ["Cookie consent (GDPR)", "Planned", "Consent banner, conditional analytics loading"],
            ["Staff Only Mode OFF", "Pending", "Single toggle flip in Admin > System Settings when ready"],
        ]
    )

    # 12. Deployment Status
    doc.add_heading("12. Deployment Status", level=1)
    add_table_from_data(doc,
        ["Environment", "Status", "URL", "Database"],
        [
            ["Production", "Staff Only Mode (locked)", "rent-a-vacation.com", "Supabase PROD"],
            ["Staging/Preview", "Active development", "Vercel preview URLs", "Supabase DEV"],
        ]
    )
    doc.add_paragraph()
    for b in [
        "21 migrations deployed to both DEV and PROD",
        "17 edge functions deployed to PROD (seed-manager DEV-only by design)",
        "CI/CD: GitHub Actions on push to main and PRs targeting main",
        "Secrets configured: RESEND_API_KEY, STRIPE_SECRET_KEY, NEWSAPI_KEY, OPENROUTER_API_KEY (both environments)",
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)

    # 13. Performance Metrics
    doc.add_heading("13. Performance Metrics — INDUSTRY DATA + PROJECTED", level=1)
    add_table_from_data(doc,
        ["Metric", "Value", "Label"],
        [
            ["Voice Search Adoption", "34% of all searches", "PROJECTED"],
            ["Voice Search Success Rate", "87%", "PROJECTED"],
            ["Voice Search NPS", "+68", "PROJECTED"],
            ["Voice vs Manual Conversion Boost", "+23%", "PROJECTED"],
            ["Listing Completion Time", "8 min (was 22 min, -64%)", "PROJECTED"],
            ["Listing Completion Rate", "94% (was 67%, +27%)", "PROJECTED"],
            ["Owner Satisfaction", "4.7 stars (was 3.8, +0.9)", "PROJECTED"],
            ["Resort Coverage", "117 resorts, 351 unit types, 10+ countries", "BUILT"],
            ["Automated Test Count", "306", "BUILT"],
        ]
    )
    add_blockquote(doc, "Honesty Framework: BUILT = deployed and demonstrable in the codebase. INDUSTRY DATA = published research from third-party sources. PROJECTED = forward-looking estimates based on industry benchmarks and internal modeling. Never present projections as actuals.")

    # Glossary
    doc.add_heading("Glossary", level=1)
    add_body(doc, "All branded terms below are RAV-coined names \u2014 proprietary marketing terms created by Rent-A-Vacation. They are not industry-standard terms.", italic=True, size=9)

    doc.add_heading("RAV-Coined Terms", level=2)
    add_table_from_data(doc,
        ["Term", "Definition"],
        [
            ["RAV", "Short for Rent-A-Vacation. Used in informal contexts, internal docs, and UI where space is limited"],
            ["RAVIO", "Rent-A-Vacation Intelligent Operator. The AI assistant brand identity for both voice (Ask RAVIO) and text chat (Chat with RAVIO)"],
            ["Name Your Price", "The bidding feature \u2014 travelers submit their own price offer on any listing where the owner has opted in"],
            ["Vacation Wishes", "The travel request feature \u2014 reverse auction where travelers post their dream trip and owners compete with proposals"],
            ["RAV SmartPrice", "Fair value scoring system using P25-P75 percentile analysis of comparable accepted bids"],
            ["Fee Freedom Calculator", "Public break-even calculator showing owners how many weeks to rent to cover maintenance fees"],
            ["TrustShield", "Owner verification program with progressive trust levels (New \u2192 Verified \u2192 Trusted \u2192 Premium). Includes document upload and admin review"],
            ["PaySafe", "Escrow payment system \u2014 holds traveler funds from booking until check-in is confirmed. Owners receive payout after checkout + 5 days"],
            ["ResortIQ", "Curated database of 117 resorts and 351 unit types from 9 vacation club brands. Auto-populates listing specs"],
            ["RAV Command", "Executive dashboard with proprietary metrics, market data integrations, and live industry feed. For RAV leadership only"],
            ["Owner's Edge", "Owner dashboard suite with 6 business intelligence sections: earnings, pricing, bids, listings, fee tracking"],
            ["Liquidity Score", "Proprietary marketplace health metric measuring supply-demand matching efficiency"],
            ["Bid Spread Index", "Proprietary price discovery metric measuring how closely bids track listed prices"],
            ["Demand Signals", "Real-time indicators showing owners matching travel request count and max budget while creating listings"],
        ]
    )

    doc.add_heading("Industry Terms", level=2)
    add_table_from_data(doc,
        ["Term", "Definition"],
        [
            ["Timeshare / Vacation Ownership", "Property ownership model where multiple buyers share rights to use a vacation property, typically in one-week intervals"],
            ["Maintenance Fees", "Annual fees charged by vacation clubs for property upkeep, regardless of whether the owner uses their allotted time"],
            ["Escrow", "Financial arrangement where a third party holds funds until conditions are met"],
            ["Per-Night Rate", "Industry-standard pricing model charging per night of stay (vs lump-sum per-week)"],
            ["P2P Marketplace", "Peer-to-peer marketplace connecting individual sellers directly with buyers"],
            ["BYOK", "Bring Your Own Key \u2014 users supply their own API keys for third-party data integrations"],
            ["RLS", "Row Level Security \u2014 PostgreSQL feature restricting database access based on user identity at the row level"],
        ]
    )

    # Footer
    add_footer(doc, "Prepared for RAV Partners \u2014 Confidential \u2014 Draft\nGenerated February 22, 2026. All statistics verified against source code and database schema.\nRent-A-Vacation | rent-a-vacation.com | Name Your Price. Book Your Paradise.\nQuestions: support@rent-a-vacation.com")

    output_path = os.path.join(SCRIPT_DIR, "RAV-roadmap-draft-02222026.docx")
    doc.save(output_path)
    print(f"Roadmap saved: {output_path}")
    return output_path


# ============================================================
# STATUS REPORT DOCUMENT
# ============================================================

def generate_status_report():
    doc = create_branded_doc("Status Report")
    add_logo_header(doc, doc_title="Development Status Report")
    add_page_numbers(doc)

    # Metadata
    add_metadata(doc, [
        ("Date", "February 22, 2026"),
        ("Prepared by", "Sujit (RAV Owner / Lead Developer)"),
        ("Version", "v0.9.0"),
        ("Platform Status", "Pre-Launch (Staff Only Mode \u2014 deployed to production, locked for internal testing)"),
        ("Last Updated", "February 22, 2026 at 11:30 PM EST"),
    ])
    doc.add_paragraph()

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    add_body(doc, "Rent-A-Vacation (RAV) is a peer-to-peer vacation rental marketplace for timeshare and vacation club owners. The platform is feature-complete for MVP with 19 completed development phases, covering the full owner-to-traveler lifecycle: property registration, listing management, AI-powered search, bidding/negotiation, Stripe payments, escrow, owner confirmation, check-in verification, and payout processing.")
    add_body(doc, "All code is deployed to production and currently locked behind \"Staff Only Mode\" for pre-launch testing and seed data validation.")
    doc.add_paragraph()
    add_body(doc, "Platform Health Dashboard", bold=True, size=12, color=DEEP_TEAL)
    add_table_from_data(doc,
        ["Metric", "Value", "Status"],
        [
            ["Automated Tests", "306 (all passing)", "✅"],
            ["TypeScript Errors", "0", "✅"],
            ["ESLint Errors", "0", "✅"],
            ["Production Build", "Clean", "✅"],
            ["Database Migrations", "21 (deployed to DEV + PROD)", "✅"],
            ["Edge Functions", "17 (deployed to PROD)", "✅"],
            ["Completed Phases", "19 + supplementary tracks", "✅"],
        ]
    )

    # 2. Technology Stack
    doc.add_heading("2. Technology Stack", level=1)
    add_table_from_data(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React 18 + TypeScript + Vite + SWC", "Single-page application with strict typing"],
            ["Styling", "Tailwind CSS + shadcn/ui (Radix primitives)", "Utility-first CSS with accessible component library"],
            ["Routing", "React Router v6", "Client-side routing with protected routes"],
            ["Data Fetching", "TanStack React Query v5", "Server state management, caching, optimistic updates"],
            ["Forms", "React Hook Form + Zod", "Schema-validated forms"],
            ["Auth", "Supabase Auth", "Email/password, Google OAuth, admin-approved signups"],
            ["Database", "Supabase PostgreSQL", "Row Level Security (RLS), pg_cron, pg_net"],
            ["Backend", "Supabase Edge Functions (Deno)", "17 serverless functions"],
            ["Payments", "Stripe Checkout", "Payment capture, escrow hold, webhooks"],
            ["Email", "Resend API", "Branded transactional emails from notifications@updates.rent-a-vacation.com"],
            ["Voice AI", "VAPI + Deepgram Nova-3", "Voice transcription and natural language property search"],
            ["Text AI", "OpenRouter (Gemini 3 Flash)", "LLM chat with SSE streaming and tool calling"],
            ["Charts", "Recharts", "Dashboard analytics and data visualization"],
            ["Hosting", "Vercel (frontend) + Supabase (backend)", "Auto-deploy from GitHub on merge to main"],
            ["CI/CD", "GitHub Actions", "5-job pipeline: lint, typecheck, unit tests, E2E, Percy visual regression"],
            ["PWA", "vite-plugin-pwa + Workbox", "Service worker (59 precached entries), install prompt, offline detection"],
        ]
    )

    # 3. Feature Inventory
    doc.add_heading("3. Feature Inventory", level=1)

    # 3.1
    doc.add_heading("3.1 Core Marketplace Features", level=2)
    add_table_from_data(doc,
        ["Feature", "Description", "Database Tables"],
        [
            ["User Registration", "Email/password + Google OAuth with admin approval workflow", "profiles, user_roles"],
            ["Role-Based Access (RBAC)", "5 roles: RAV Owner, RAV Admin, RAV Staff, Property Owner, Renter", "user_roles (enum: app_role)"],
            ["Property Registration", "Multi-step form with resort search (117 resorts), auto-populate specs, image upload", "properties, property-images bucket"],
            ["Listing Management", "Draft → Pending Approval → Active lifecycle, per-night pricing, 4 cancellation policies", "listings (nightly_rate, owner_price, rav_markup, final_price)"],
            ["Booking Flow", "Browse → View → Book Now → Stripe Checkout → Payment → Confirmation", "bookings, booking_confirmations"],
            ["Escrow (PaySafe)", "Funds held until check-in confirmed. Released to owner after checkout + 5 days", "booking_confirmations (escrow_status)"],
            ["Owner Confirmation Timer", "Configurable countdown (default 60 min), up to 2 × 30-min extensions, auto-cancel", "booking_confirmations"],
            ["Check-in Verification", "Traveler confirms arrival or reports issues (access, safety, mismatch)", "checkin_confirmations"],
            ["Cancellation System", "4 policies: Flexible, Moderate, Strict, Super Strict. Refund calculation engine", "cancellation_requests"],
            ["Owner Verification (TrustShield)", "Document upload, 4 trust levels, admin review workflow", "owner_verifications, verification_documents"],
        ]
    )

    # 3.2
    doc.add_heading("3.2 AI-Powered Search", level=2)
    add_table_from_data(doc,
        ["Feature", "Name", "Technology", "Key Details"],
        [
            ["Voice Search", "Ask RAVIO", "VAPI + Deepgram Nova-3", "Natural language queries, 300ms endpointing, smart denoising, LiveKit smart endpointing, keyword boosts. Shared property-search.ts module with state name expansion"],
            ["Text Chat", "Chat with RAVIO", "OpenRouter (Gemini 3 Flash)", "SSE streaming, tool calling (search_properties), 4 context-aware system prompts (rentals, property-detail, bidding, general). JWT auth, 60 req/min rate limit"],
            ["Resort Database", "ResortIQ", "PostgreSQL", "117 resorts (Hilton 62, Marriott 40, Disney 15), 351 unit types, 10+ countries. Auto-populate bedrooms, bathrooms, max guests, square footage"],
        ]
    )

    # 3.3
    doc.add_heading("3.3 Bidding & Negotiation", level=2)
    add_table_from_data(doc,
        ["Feature", "Name", "Description"],
        [
            ["Standard Bids", "Name Your Price", "Travelers bid on listings where owner opted in. Owner reviews → accept/reject/counter"],
            ["Date Proposals", "—", "Bid with different dates; amount auto-computes from nightly_rate × proposed nights. Blue badge in owner's bid manager"],
            ["Travel Requests", "Vacation Wishes", "Reverse auction: travelers post destination + dates + budget, owners respond with proposals. Auto-matching on listing approval"],
            ["Inspired Requests", "—", "\"Request Similar Dates\" button on PropertyDetail pre-fills travel request. Optional owner targeting"],
            ["Demand Signals", "—", "Owners see matching travel request count + max budget while creating listings (500ms debounce)"],
            ["Auto-Matching", "—", "match-travel-requests edge function runs on listing approval, matches by destination, dates (±30 days), bedrooms, budget, brand"],
        ]
    )

    # 3.4
    doc.add_heading("3.4 Business Intelligence", level=2)
    add_table_from_data(doc,
        ["Dashboard", "Name", "Audience", "Sections"],
        [
            ["Executive", "RAV Command", "RAV Owner", "(1) Headline KPI bar, (2) Business Performance (4 charts), (3) Marketplace Health (Liquidity Score gauge, supply/demand map, voice funnel), (4) Market Intelligence (AirDNA + STR via BYOK), (5) Industry Feed, (6) Unit Economics"],
            ["Owner", "Owner's Edge", "Property Owners", "(1) Headline Stats, (2) Earnings Timeline (AreaChart + fee target), (3) My Listings Table (status/Fair Value badges), (4) Bid Activity Feed, (5) Pricing Intelligence, (6) Maintenance Fee Tracker"],
            ["Fair Value", "RAV SmartPrice", "All users", "P25-P75 percentile analysis. Tiers: below_market, fair_value, above_market. Role-specific messaging"],
            ["Calculator", "Fee Freedom Calculator", "Public", "Break-even analysis for 9 brands, 4 unit types. Progress bars + CTA"],
        ]
    )

    # 3.5
    doc.add_heading("3.5 Admin & Operations", level=2)
    add_table_from_data(doc,
        ["Capability", "Details"],
        [
            ["Admin Dashboard", "12 tabs: Overview, Users, Listings, Bookings, Properties, Verifications, Escrow, Payouts, Financials, Issues, Voice, Memberships"],
            ["Voice Admin", "5 sections: Config info, Tier quota manager, Per-user overrides, Usage dashboard (charts + top users), Observability (log viewer + alert thresholds)"],
            ["Staff Only Mode", "Pre-launch lock with 3-layer enforcement: (1) Database RLS, (2) Login signs out non-RAV users, (3) Signup shows \"Coming Soon\". Toggle in Admin > System Settings"],
            ["Seed Data System", "DEV-only 3-layer system with production guard. Layer 1: 8 foundation users. Layer 2: 10 properties, 30 listings. Layer 3: 50 renters, 110 bookings, 20 bids. Password: SeedTest2026!"],
        ]
    )

    # 4. Membership & Pricing
    doc.add_heading("4. Membership & Pricing", level=1)
    doc.add_heading("4.1 Membership Tiers (6 total)", level=2)
    add_body(doc, "Renter Tiers:", bold=True)
    add_table_from_data(doc,
        ["Tier", "Monthly Price", "Voice Searches/Day", "Benefits"],
        [
            ["Free", "$0", "5", "Browse listings, place bids, post travel requests"],
            ["Plus", "$9.99", "25", "Priority support, saved searches"],
            ["Premium", "$24.99", "Unlimited", "Early access, concierge service"],
        ]
    )
    doc.add_paragraph()
    add_body(doc, "Owner Tiers:", bold=True)
    add_table_from_data(doc,
        ["Tier", "Monthly Price", "Commission Rate", "Benefits"],
        [
            ["Free", "$0", "15% (default)", "List properties, basic dashboard, bid management"],
            ["Pro", "$19.99", "13% (−2%)", "Analytics, priority listing placement"],
            ["Business", "$49.99", "10% (−5%)", "Multi-property management, API access, dedicated support"],
        ]
    )
    add_blockquote(doc, "Source: Migration 011 (membership_tiers table). The base commission rate (currently 15%) is admin-configurable in Admin > System Settings (platform_commission_rate). Stripe processing fees (~2.9%) are absorbed by RAV within the service fee margin.")

    # 4.2 Brands
    doc.add_heading("4.2 Supported Vacation Club Brands (9)", level=2)
    brands = ["Hilton Grand Vacations (62 resorts in ResortIQ)", "Marriott Vacation Club (40 resorts)", "Disney Vacation Club (15 resorts)", "Wyndham Destinations", "Hyatt Residence Club", "Bluegreen Vacations", "Holiday Inn Club Vacations", "WorldMark by Wyndham", "Other / Independent Resort"]
    for i, b in enumerate(brands, 1):
        p = doc.add_paragraph(f"{i}. {b}")
        for run in p.runs:
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)
    add_blockquote(doc, "Source: VACATION_CLUB_BRANDS in calculatorLogic.ts and vacation_club_brand database enum.")

    # 5. Edge Functions
    doc.add_heading("5. Edge Functions (17 total)", level=1)
    add_table_from_data(doc,
        ["#", "Function", "Trigger", "Purpose"],
        [
            ["1", "create-booking-checkout", "Client call", "Creates Stripe Checkout session with tier-aware commission"],
            ["2", "verify-booking-payment", "Stripe webhook", "Validates payment, creates booking + confirmation with acceptance timer, sends emails"],
            ["3", "send-email", "Client call", "Generic transactional email via Resend API"],
            ["4", "send-approval-email", "Client call", "Approval/rejection notifications (4 template variants)"],
            ["5", "send-booking-confirmation-reminder", "Client/internal", "Owner deadline reminders + acceptance notifications"],
            ["6", "send-cancellation-email", "Internal", "Cancellation status notifications (4 variants)"],
            ["7", "send-contact-form", "Client call", "Contact form submission with confirmation"],
            ["8", "send-verification-notification", "Client call", "Admin notification on doc upload"],
            ["9", "process-deadline-reminders", "CRON (30 min)", "Scan deadlines, reminders, timeouts, travel request expiry warnings"],
            ["10", "match-travel-requests", "Internal (admin)", "Auto-match listings to travel requests (budget-aware, deduped)"],
            ["11", "voice-search", "VAPI webhook", "Property search via voice, shared search module"],
            ["12", "text-chat", "Client call", "OpenRouter LLM, SSE streaming, tool calling, 4 modes"],
            ["13", "seed-manager", "Client call", "DEV-only 3-layer seed data (production-guarded)"],
            ["14", "fetch-industry-news", "Client call", "NewsAPI + Google News RSS (60-min cache)"],
            ["15", "fetch-macro-indicators", "Client call", "FRED consumer confidence + travel data"],
            ["16", "fetch-airdna-data", "Client call", "AirDNA market comparisons (BYOK)"],
            ["17", "fetch-str-data", "Client call", "STR hospitality benchmarks (BYOK)"],
        ]
    )

    doc.add_heading("Required Secrets (Supabase Dashboard)", level=2)
    add_table_from_data(doc,
        ["Secret", "Used by", "Environments"],
        [
            ["RESEND_API_KEY", "All email functions", "DEV + PROD"],
            ["STRIPE_SECRET_KEY", "create-booking-checkout, verify-booking-payment", "DEV + PROD"],
            ["NEWSAPI_KEY", "fetch-industry-news", "DEV + PROD"],
            ["OPENROUTER_API_KEY", "text-chat", "DEV + PROD"],
            ["IS_DEV_ENVIRONMENT", "seed-manager (production guard)", "DEV only"],
        ]
    )

    # 6. Email System
    doc.add_heading("6. Email System", level=1)
    add_body(doc, "17 transactional email types via Resend API, using branded HTML templates from _shared/email-template.ts.", bold=True)
    add_table_from_data(doc,
        ["Category", "Email", "Trigger", "Recipient"],
        [
            ["Account", "Welcome", "User signup", "New user"],
            ["", "User Approved", "Admin approves", "User"],
            ["", "User Rejected", "Admin rejects", "User"],
            ["Listings", "Listing Approved", "Admin approves", "Owner"],
            ["", "Listing Rejected", "Admin rejects", "Owner"],
            ["", "Listing Submitted", "Owner submits", "RAV admin"],
            ["Bookings", "Booking Confirmed", "Payment verified", "Traveler"],
            ["", "Check-in Reminder", "CRON, near arrival", "Traveler"],
            ["Owner Confirm.", "Confirmation Request", "Payment verified", "Owner"],
            ["", "Extension Notification", "Owner requests extension", "Renter"],
            ["", "Confirmation Timeout", "Owner times out", "Owner + Renter"],
            ["Cancellation", "Submitted", "Request created", "Traveler"],
            ["", "Approved", "Owner approves", "Traveler"],
            ["", "Denied", "Owner denies", "Traveler"],
            ["", "Counter-Offer", "Owner counter-offers", "Traveler"],
            ["Verification", "Document Uploaded", "Doc upload", "RAV admin"],
            ["Support", "Contact Form", "Form submission", "support@"],
        ]
    )

    # 7. Recent Development Activity
    doc.add_heading("7. Recent Development Activity (Sessions 14–16)", level=1)

    doc.add_heading("Session 14: Phase 19 — Flexible Date Booking + Per-Night Pricing (Feb 22)", level=2)
    bullets_s14 = [
        "Migration 020: Added nightly_rate column to listings (backfilled from owner_price / nights), requested_check_in/out on listing_bids, source_listing_id + target_owner_only on travel_requests",
        "Shared pricing utility: src/lib/pricing.ts — calculateNights() + computeListingPricing() replacing 4 duplicated functions",
        "BidFormDialog dual-mode: Standard bid vs date-proposal with auto-computed amounts",
        "InspiredTravelRequestDialog: \"Request Similar Dates\" from listing detail, pre-fills form, optional owner targeting",
        "Owner listing form: Switched from lump-sum \"Your Asking Price\" to \"Nightly Rate\" with live price breakdown",
        "16 new tests (289 total). PR #20 merged, migration deployed to DEV + PROD",
    ]
    for b in bullets_s14:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)

    doc.add_heading("Session 15: Content Accuracy Audit (Feb 22)", level=2)
    for b in [
        "Fixed commission rate (10% → 15%) across 7 code files + 3 test files",
        "Corrected brand list (Westgate → WorldMark, 8 → 9 brands)",
        "Fixed voice quota display (flat 10/day → tier-based from database)",
        "Added 9 missing sections to Documentation.tsx admin manual",
        "Established Content Accuracy (MANDATORY) policy in CLAUDE.md",
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)

    doc.add_heading("Session 16: Voice Tracks C-D — Admin Controls + Observability (Feb 22)", level=2)
    for b in [
        "Migration 021: voice_search_logs table, voice_user_overrides table, 3 RPCs, 2 alert threshold settings",
        "Admin Dashboard \"Voice\" tab: 5 sections — config info, tier quota manager, per-user overrides, usage dashboard, observability",
        "Auto-logging: All voice searches automatically logged with query, results count, duration, success status",
        "17 new tests (306 total)",
    ]:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = BRAND_FONT
            run.font.size = Pt(10)

    # 8. Deployment Status
    doc.add_heading("8. Deployment Status", level=1)
    add_table_from_data(doc,
        ["Environment", "Status", "URL", "Database"],
        [
            ["Production", "Staff Only Mode (locked)", "rent-a-vacation.com", "Supabase PROD"],
            ["Staging/Preview", "Active development", "Vercel preview URLs", "Supabase DEV"],
        ]
    )
    doc.add_paragraph()
    add_table_from_data(doc,
        ["Resource", "Count", "Deployment"],
        [
            ["Database Migrations", "21", "Both DEV + PROD"],
            ["Edge Functions", "17", "PROD (seed-manager DEV-only)"],
            ["Automated Tests", "306", "All passing"],
            ["GitHub PRs Merged", "#12–#21", "All to main"],
        ]
    )

    # 9. Next Priorities
    doc.add_heading("9. Next Priorities", level=1)
    add_table_from_data(doc,
        ["Priority", "Phase", "Description", "Timeline"],
        [
            ["1", "Phase 20A-C", "Accounting, Tax & Fee Framework (fee breakdown, Stripe Tax, reporting)", "Pre-launch"],
            ["2", "Phase 3", "Voice Everywhere (voice-assisted listing, booking, bidding)", "Q2 2026"],
            ["3", "Phase 12", "Native App Shells via Capacitor (Android + iOS)", "Q2-Q3 2026"],
            ["4", "Phase 20D-F", "QuickBooks integration, 1099-K compliance, automated tax filing", "Post-launch"],
            ["5", "Phase 21", "Partial-Week Booking (listing splits, minimum stay)", "When demand validates"],
            ["6", "Phase 6", "Advanced Features (saved searches, map view, calendar integration)", "Q3 2026"],
        ]
    )

    # 10. Performance Metrics
    doc.add_heading("10. Performance Metrics", level=1)
    add_table_from_data(doc,
        ["Metric", "Value", "Label"],
        [
            ["Resort Coverage", "117 resorts, 351 unit types, 10+ countries", "BUILT"],
            ["Automated Test Count", "306", "BUILT"],
            ["Voice Search Adoption", "34% of all searches", "PROJECTED"],
            ["Voice Search Success Rate", "87%", "PROJECTED"],
            ["Voice NPS", "+68", "PROJECTED"],
            ["Listing Completion Time", "8 min (was 22 min, −64%)", "PROJECTED"],
            ["Listing Completion Rate", "94% (was 67%, +27%)", "PROJECTED"],
            ["Owner Satisfaction", "4.7 stars (was 3.8, +0.9)", "PROJECTED"],
        ]
    )
    add_blockquote(doc, "Honesty Framework: BUILT = deployed and demonstrable in the codebase. INDUSTRY DATA = published research from third-party sources. PROJECTED = forward-looking estimates based on industry benchmarks and internal modeling.")

    # Footer
    add_footer(doc, "Generated February 22, 2026. All statistics verified against source code and database schema.\nRent-A-Vacation | rent-a-vacation.com | Name Your Price. Book Your Paradise.")

    output_path = os.path.join(SCRIPT_DIR, "RAV-Development-Status-Report-02222026.docx")
    doc.save(output_path)
    print(f"Status Report saved: {output_path}")
    return output_path


if __name__ == "__main__":
    print("Generating branded .docx files...")
    generate_roadmap()
    generate_status_report()
    print("Done!")
